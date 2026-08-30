#!/usr/bin/env python3
import argparse
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from common import read_json

BLUE="1F4E78"; LIGHT="D9EAF7"; ALERT_FILL="FFF2CC"; ALERT_TEXT="7F6000"
def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement('w:tblHeader'); e.set(qn('w:val'),'true'); trPr.append(e)
def set_cell(cell, text, bold=False, color=None, size=8.5):
    cell.text=""; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(str(text)); r.bold=bold; r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def setup(doc):
    sec=doc.sections[0]; sec.top_margin=Cm(2.2); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.4); sec.right_margin=Cm(2.2)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Microsoft YaHei'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); normal.font.size=Pt(10.5)
    normal.paragraph_format.line_spacing=1.5; normal.paragraph_format.space_after=Pt(6)
    for name,size,color in [('Title',22,BLUE),('Heading 1',16,BLUE),('Heading 2',13,BLUE)]:
        s=styles[name]; s.font.name='Microsoft YaHei'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color)
def title_page(doc, project, year, subtitle=None, period_label=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(130)
    r=p.add_run(project); r.bold=True; r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); r.font.size=Pt(24); r.font.color.rgb=RGBColor.from_string(BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(subtitle or '生产运营数据分析报告'); r.bold=True; r.font.size=Pt(20); r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(45); p.add_run(period_label or f'{year} 年度数据分析').font.size=Pt(13)
    doc.add_page_break()
def add_note(doc, text):
    t=doc.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,LIGHT); set_cell(c,text,size=9); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
def add_critical_input(doc, items):
    if not items: return
    t=doc.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,ALERT_FILL)
    c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=p.add_run('关键输入提醒'); r.bold=True; r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string(ALERT_TEXT)
    for index,item in enumerate(items,1):
        p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.left_indent=Cm(.5); p.paragraph_format.first_line_indent=Cm(-.5)
        text=item.get('text') or item.get('message') or str(item); r=p.add_run(f'{index}. {text}'); r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(ALERT_TEXT)
def add_monthly(doc, model):
    sec=doc.add_section(); sec.orientation=WD_ORIENT.LANDSCAPE; sec.page_width,sec.page_height=sec.page_height,sec.page_width; sec.left_margin=Cm(1.2); sec.right_margin=Cm(1.2); sec.top_margin=Cm(1.2); sec.bottom_margin=Cm(1.2)
    doc.add_heading('月度进出水指标统计', level=2)
    periods=[p[5:]+'月' for p in model['monthly_table']['periods']]; table=doc.add_table(rows=1, cols=len(periods)+1); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    headers=['指标']+periods
    table.rows[0].height=Cm(.75); table.rows[0].height_rule=WD_ROW_HEIGHT_RULE.EXACTLY
    for i,x in enumerate(headers): set_cell(table.rows[0].cells[i],x,True,'FFFFFF',7); shade(table.rows[0].cells[i],BLUE)
    repeat_header(table.rows[0])
    for row in model['monthly_table']['rows']:
        tr=table.add_row(); tr.height=Cm(.72); tr.height_rule=WD_ROW_HEIGHT_RULE.EXACTLY; cells=tr.cells; set_cell(cells[0],row['label'].replace('（mg/L）','').replace('（万 m³/d）',''),True,size=6.5); shade(cells[0],'EAF2F8')
        for i,v in enumerate(row['values']): set_cell(cells[i+1],'' if v is None else f"{v:.{row['decimals']}f}",size=6.5)
def portrait_section(doc):
    sec=doc.add_section(); sec.orientation=WD_ORIENT.PORTRAIT; sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.2); sec.right_margin=Cm(2.2)
def add_removal(doc, model):
    doc.add_heading('污染物去除效率统计', level=2)
    if not model['removal_table']:
        add_note(doc,'源文件未提供进出水水质、污染物去除量和排放量，本次不计算污染物去除效率，也不评价达标性。'); return
    table=doc.add_table(1,4); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,x in enumerate(['指标','去除量（kg）','排放量（kg）','质量平衡去除率']): set_cell(table.rows[0].cells[i],x,True,'FFFFFF'); shade(table.rows[0].cells[i],BLUE)
    repeat_header(table.rows[0])
    for x in model['removal_table']:
        cells=table.add_row().cells
        for i,v in enumerate([x['pollutant'],f"{x['removed_kg']:,.2f}",f"{x['emitted_kg']:,.2f}",f"{x['rate_pct']:.2f}%"]): set_cell(cells[i],v)
def build(model,draft,charts,out,chapter=False):
    doc=Document(); setup(doc)
    if not chapter: title_page(doc,model['project'],model['analysis_year'],period_label=model.get('period_label'))
    else: doc.add_heading('运营数据现状评估及建设需求',0)
    section_map={s['section_id']:s for s in draft['sections']}
    for idx,definition in enumerate(model['sections'],1):
        sid=definition['section_id']; drafted_section=section_map.get(sid,{})
        if sid=='efficiency' and not model.get('removal_table') and not drafted_section.get('paragraphs'):
            continue
        doc.add_heading((f'{idx} ' if not chapter else '')+definition['title'], level=1)
        for para in section_map.get(sid,{}).get('paragraphs',[]):
            p=doc.add_paragraph(para['text']); p.paragraph_format.first_line_indent=Cm(0.74)
        callouts=section_map.get(sid,{}).get('callouts',[])
        if callouts: add_critical_input(doc,callouts)
        if sid=='flow_quality':
            add_monthly(doc,model); doc.add_picture(str(charts/'monthly-trends.png'),width=Cm(24)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; portrait_section(doc)
        elif sid=='efficiency': add_removal(doc,model)
        elif sid=='cost':
            if model['cost_table']:
                table=doc.add_table(1,3); table.style='Table Grid'
                for i,x in enumerate(['成本类别','年度金额（元）','本口径占比']): set_cell(table.rows[0].cells[i],x,True,'FFFFFF'); shade(table.rows[0].cells[i],BLUE)
                for x in model['cost_table']:
                    c=table.add_row().cells
                    for i,v in enumerate([x['label'],f"{x['cost_cny']:,.2f}",f"{x['share_pct']:.2f}%"]): set_cell(c[i],v)
                doc.add_picture(str(charts/'cost-share.png'),width=Cm(15)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
            else:
                pass
            if model.get('resource_table'):
                doc.add_heading('主要资源消耗同比',level=2); table=doc.add_table(1,4); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
                for i,x in enumerate(['指标','2025年同期','2026年同期','同比变化']): set_cell(table.rows[0].cells[i],x,True,'FFFFFF'); shade(table.rows[0].cells[i],BLUE)
                repeat_header(table.rows[0])
                for x in model['resource_table']:
                    c=table.add_row().cells; change=x.get('change_pct'); values=[x['label'],f"{x['value_2025']:,.2f} {x['unit']}",f"{x['value_2026']:,.2f} {x['unit']}",'' if change is None else f"{change:+.2f}%"]
                    for i,v in enumerate(values): set_cell(c[i],v)
                if (charts/'yoy-change.png').exists(): doc.add_picture(str(charts/'yoy-change.png'),width=Cm(16)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif sid=='limitations':
            add_critical_input(doc,model.get('critical_inputs',[]))
            if model.get('review_only_issues'):
                add_note(doc,f"其余 {len(model['review_only_issues'])} 项数据质量事项不进入正式结论，详细来源、差异值和处理状态见《待确认事项.xlsx》。")
    Path(out).parent.mkdir(parents=True,exist_ok=True); doc.save(out)
def main():
    p=argparse.ArgumentParser(); p.add_argument('--model',required=True); p.add_argument('--draft',required=True); p.add_argument('--charts',required=True); p.add_argument('--standalone'); p.add_argument('--chapter')
    a=p.parse_args(); model,draft=read_json(a.model),read_json(a.draft); charts=Path(a.charts)
    if a.standalone: build(model,draft,charts,a.standalone,False)
    if a.chapter: build(model,draft,charts,a.chapter,True)
if __name__=='__main__': main()

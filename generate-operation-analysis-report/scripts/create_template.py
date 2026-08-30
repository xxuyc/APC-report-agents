#!/usr/bin/env python3
import argparse
from pathlib import Path
from docx import Document
from build_docx import setup
def main():
    p=argparse.ArgumentParser(); p.add_argument('--output',required=True); a=p.parse_args()
    d=Document(); setup(d); d.add_heading('生产运营数据分析报告',0); d.add_paragraph('此文件为 Harness V0.1 通用版式资产，正式内容由流水线生成。')
    Path(a.output).parent.mkdir(parents=True,exist_ok=True); d.save(a.output)
if __name__=='__main__': main()

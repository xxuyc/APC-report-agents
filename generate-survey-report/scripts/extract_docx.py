#!/usr/bin/env python3
"""Extract a review-friendly outline from a DOCX without changing it."""

import json
import sys
from pathlib import Path

from docx import Document


def text_of(cell):
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: extract_docx.py INPUT.docx OUTPUT.json")

    source = Path(sys.argv[1])
    document = Document(source)
    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            paragraphs.append({
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "text": text,
            })

    tables = []
    for index, table in enumerate(document.tables, start=1):
        rows = [[text_of(cell) for cell in row.cells] for row in table.rows]
        tables.append({"index": index, "rows": rows})

    payload = {
        "source_file": str(source),
        "paragraphs": paragraphs,
        "tables": tables,
        "section_count": len(document.sections),
    }
    Path(sys.argv[2]).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()

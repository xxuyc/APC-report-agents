#!/usr/bin/env python3
"""Run deterministic checks that do not require rendering Word pages."""

import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: check_docx.py REPORT.docx OUTPUT.json")
    document = Document(sys.argv[1])
    errors, warnings = [], []
    texts = [paragraph.text.strip() for paragraph in document.paragraphs]
    furniture_texts = []
    for section in document.sections:
        for part in (section.header, section.first_page_header, section.even_page_header,
                     section.footer, section.first_page_footer, section.even_page_footer):
            furniture_texts.extend(paragraph.text.strip() for paragraph in part.paragraphs)
    placeholders = [text for text in texts + furniture_texts if re.search(r"【[^】]+】", text)]
    if placeholders:
        errors.append({"type": "unresolved_placeholder", "count": len(placeholders),
                       "examples": placeholders[:5]})
    full_text = "\n".join(texts)
    for phrase in ("调研表显示", "现场填写信息表明", "调研资料显示", "调研表记录"):
        if phrase in full_text:
            errors.append({"type": "mechanical_source_preface", "phrase": phrase})
    if "12期" in full_text or "3期" in full_text:
        errors.append({"type": "nonstandard_phase_name", "message": "存在未规范的阶段名称。"})
    with zipfile.ZipFile(sys.argv[1]) as archive:
        names = set(archive.namelist())
        if "word/comments.xml" in names:
            errors.append({"type": "unresolved_comments", "message": "清洁稿不得包含 Word 批注。"})
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        if "<w:ins" in document_xml or "<w:del" in document_xml:
            errors.append({"type": "tracked_changes_present", "message": "清洁稿不得包含修订痕迹。"})
    headings = [(index, paragraph.text.strip()) for index, paragraph in enumerate(document.paragraphs)
                if paragraph.style and paragraph.style.name.startswith("Heading")]
    for position, (index, title) in enumerate(headings):
        next_index = headings[position + 1][0] if position + 1 < len(headings) else len(document.paragraphs)
        body = [text for text in texts[index + 1:next_index] if text]
        if not title:
            errors.append({"type": "empty_heading", "paragraph": index})
        elif not body:
            warnings.append({"type": "empty_section", "title": title})
    if not headings:
        errors.append({"type": "missing_headings", "message": "报告中未识别到标题样式。"})
    warnings.append({"type": "visual_qa_required",
                     "message": "表格跨页、图片、空白页、目录和页码需由 Agent 渲染后确认。"})
    result = {"valid": not errors, "counts": {"error": len(errors), "warning": len(warnings)},
              "errors": errors, "warnings": warnings,
              "document": {"paragraphs": len(document.paragraphs), "tables": len(document.tables),
                           "headings": len(headings)}}
    Path(sys.argv[2]).write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    if errors:
        raise SystemExit(2)


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Extract auditable evidence blocks from approved supporting documents."""

import json
import sys
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def docx_blocks(path):
    document = Document(path)
    blocks = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = paragraph.text.strip()
        if text:
            blocks.append({"locator": f"paragraph:{index}", "text": text})
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                blocks.append({"locator": f"table:{table_index}:row:{row_index}", "text": text})
    return blocks


def pdf_blocks(path):
    blocks = []
    for page_index, page in enumerate(PdfReader(path).pages, 1):
        for paragraph_index, text in enumerate((page.extract_text() or "").splitlines(), 1):
            if text.strip():
                blocks.append({"locator": f"page:{page_index}:line:{paragraph_index}", "text": text.strip()})
    return blocks


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: extract_supporting.py source-index.json OUTPUT.json")
    source_index = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    root = Path(source_index["root"])
    evidence = []
    for source in source_index.get("sources", []):
        if source.get("source_type") != "supporting_document":
            continue
        path = root / source["relative_path"]
        suffix = path.suffix.lower()
        if suffix == ".docx":
            blocks = docx_blocks(path)
        elif suffix == ".pdf":
            blocks = pdf_blocks(path)
        elif suffix in (".txt", ".md"):
            blocks = [{"locator": f"line:{i}", "text": line.strip()} for i, line in enumerate(
                path.read_text(encoding="utf-8", errors="replace").splitlines(), 1) if line.strip()]
        else:
            blocks = []
        for index, block in enumerate(blocks, 1):
            evidence.append({
                "evidence_id": f"{source['source_id']}-B{index:04d}", "source_id": source["source_id"],
                "source_file": source["file_name"], "usage": source["usage"],
                "role": "background_fact", "confirmation": "confirmed" if source["usage"] == "fact_source" else "feedback_only",
                **block,
            })
    Path(sys.argv[2]).write_text(json.dumps({"schema_version": "0.2", "evidence": evidence},
                                           ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()

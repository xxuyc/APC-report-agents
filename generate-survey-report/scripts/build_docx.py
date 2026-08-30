#!/usr/bin/env python3
"""Build a reviewable Word report from validated agent-authored content."""

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_run_font(run, size=12, bold=False, font="宋体"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def clear_body(document):
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def item_text(item):
    if isinstance(item, str):
        return item
    return item.get("text") or item.get("item") or item.get("message") or item.get("question") or ""


def replace_project_placeholders(document, project_name):
    if not project_name:
        return
    for section in document.sections:
        for part in (section.header, section.first_page_header, section.even_page_header,
                     section.footer, section.first_page_footer, section.even_page_footer):
            for paragraph in part.paragraphs:
                for run in paragraph.runs:
                    run.text = run.text.replace("【项目名称】", project_name)


def add_field(paragraph, instruction, cached_text="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_element = OxmlElement("w:instrText")
    instruction_element.set(qn("xml:space"), "preserve")
    instruction_element.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = cached_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction_element, separate, text, end))
    set_run_font(run, size=9)


def configure_page_footer(document):
    for section in document.sections:
        footers = [section.footer]
        if document.settings.odd_and_even_pages_header_footer:
            footers.append(section.even_page_footer)
        for footer in footers:
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            for child in list(paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph._p.remove(child)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph.add_run("第 "), size=9)
            add_field(paragraph, "PAGE")
            set_run_font(paragraph.add_run(" 页 / 共 "), size=9)
            add_field(paragraph, "NUMPAGES")
            set_run_font(paragraph.add_run(" 页"), size=9)


def set_cell_shading(cell, color="EAF2F8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)
    margins = OxmlElement("w:tcMar")
    for edge, value in (("top", "120"), ("left", "180"), ("bottom", "120"), ("right", "180")):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)


def add_pending_callout(document, items):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(15.2)
    cell = table.cell(0, 0)
    cell.width = Cm(15.2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    set_run_font(paragraph.add_run("待确认事项"), size=10.5, bold=True, font="黑体")
    for index, item in enumerate(items, 1):
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(12)
        paragraph.paragraph_format.first_line_indent = Pt(-12)
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(f"{index}. {item}"), size=10.5)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("drafted_content")
    parser.add_argument("output")
    parser.add_argument("--template")
    parser.add_argument("--project-name", default="")
    args = parser.parse_args()

    content = json.loads(Path(args.drafted_content).read_text(encoding="utf-8"))
    document = Document(args.template) if args.template else Document()
    if args.template:
        clear_body(document)
    replace_project_placeholders(document, args.project_name)
    configure_page_footer(document)
    section = document.sections[0]
    section.top_margin = Cm(2.7)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title.paragraph_format.space_after = Pt(24)
    set_run_font(title.add_run(content.get("title", "现场调研现状评估")), size=22, bold=True, font="黑体")
    if args.project_name and args.project_name not in content.get("title", ""):
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(subtitle.add_run(args.project_name), size=15, font="黑体")
    document.add_page_break()

    for chapter in content.get("chapters", []):
        if chapter.get("status") == "omitted_not_applicable":
            continue
        heading = document.add_paragraph(style="Heading 1")
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.space_before = Pt(16)
        heading.paragraph_format.space_after = Pt(8)
        set_run_font(heading.add_run(chapter["title"]), size=16, bold=True, font="黑体")
        for item in chapter.get("paragraphs", []):
            text = item_text(item).strip()
            if not text:
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(24)
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.widow_control = True
            set_run_font(paragraph.add_run(text))
        pending = [item_text(item).strip() for item in chapter.get("to_confirm", [])]
        pending = [item for item in pending if item]
        if pending:
            add_pending_callout(document, pending)

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


if __name__ == "__main__":
    main()
